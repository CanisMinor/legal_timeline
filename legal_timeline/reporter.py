"""
reporter.py — generate a Word (.docx) report from a BranchTree.

Output document structure
-------------------------
* Title: "Key Dates Timeline"
* Subtitle: source filename + generation date
* Table: Category (left) | Date(s) & Notes (right)
  - One row per category
  - Multiple dates within a category are listed as paragraphs inside the cell
  - Unresolved relative dates are shown with a ⚠ marker and their anchor
* Appendix section: the textual branch-tree for reference
"""

from __future__ import annotations

import logging
from datetime import date, datetime
from pathlib  import Path
from typing   import Dict, List, Optional, Tuple

from docx                        import Document
from docx.shared                 import Pt, RGBColor, Inches, Cm
from docx.enum.text              import WD_ALIGN_PARAGRAPH
from docx.enum.table             import WD_ALIGN_VERTICAL
from docx.oxml.ns                import qn
from docx.oxml                   import OxmlElement

from .timeline import BranchTree

log = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Colour palette
# ---------------------------------------------------------------------------

COLOUR_HEADER_BG  = "1F3864"   # dark navy
COLOUR_HEADER_FG  = "FFFFFF"   # white
COLOUR_ROW_ALT    = "EBF2F9"   # pale blue
COLOUR_ROW_MAIN   = "FFFFFF"   # white
COLOUR_UNRESOLVED = "C00000"   # red for warnings
COLOUR_ACCENT     = "2E75B6"   # mid blue


# ---------------------------------------------------------------------------
# Low-level XML helpers (python-docx doesn't expose everything)
# ---------------------------------------------------------------------------

def _set_cell_bg(cell, hex_colour: str) -> None:
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_colour)
    tcPr.append(shd)


def _set_cell_borders(cell) -> None:
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:color"), "AAAAAA")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _para_style_normal(para) -> None:
    """Remove any built-in paragraph style that might override our formatting."""
    pPr = para._p.get_or_add_pPr()
    pStyle = OxmlElement("w:pStyle")
    pStyle.set(qn("w:val"), "Normal")
    pPr.insert(0, pStyle)


# ---------------------------------------------------------------------------
# Main reporter
# ---------------------------------------------------------------------------

class TimelineReporter:
    """
    Generate a Word document containing the key-dates table.

    Parameters
    ----------
    title:
        Document title shown at the top of the page.
    show_branch_tree:
        If ``True`` (default), append a textual branch-tree appendix.
    """

    def __init__(
        self,
        title: str = "Key Dates — Legal Timeline",
        show_branch_tree: bool = True,
    ) -> None:
        self.title            = title
        self.show_branch_tree = show_branch_tree

    # ------------------------------------------------------------------
    # Public entry point
    # ------------------------------------------------------------------

    def generate(
        self,
        tree:        BranchTree,
        output_path: str | Path,
        source_name: str = "",
    ) -> Path:
        """
        Write the report to *output_path* and return the resolved ``Path``.
        """
        output_path = Path(output_path)
        doc         = Document()

        self._set_page_layout(doc)
        self._write_title(doc, source_name)
        self._write_table(doc, tree)

        if self.show_branch_tree:
            self._write_branch_appendix(doc, tree)

        doc.save(output_path)
        log.info("Saved timeline report to %s", output_path)
        return output_path

    # ------------------------------------------------------------------
    # Page layout
    # ------------------------------------------------------------------

    def _set_page_layout(self, doc: Document) -> None:
        section = doc.sections[0]
        section.page_width    = Inches(8.5)
        section.page_height   = Inches(11)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)

    # ------------------------------------------------------------------
    # Title block
    # ------------------------------------------------------------------

    def _write_title(self, doc: Document, source_name: str) -> None:
        # Main title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(self.title)
        run.bold      = True
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

        # Subtitle
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        parts = []
        if source_name:
            parts.append(f"Source: {source_name}")
        parts.append(f"Generated: {datetime.now().strftime('%d %B %Y')}")
        sub_run = sub_para.add_run("   |   ".join(parts))
        sub_run.font.size  = Pt(10)
        sub_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
        sub_run.italic = True

        doc.add_paragraph()  # spacer

    # ------------------------------------------------------------------
    # Main table
    # ------------------------------------------------------------------

    def _write_table(self, doc: Document, tree: BranchTree) -> None:
        flat: Dict[str, List[Tuple]] = tree.flat_view()

        if not flat:
            doc.add_paragraph("No dates were extracted from this document.")
            return

        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"

        # --- Header row ---
        hdr_cells = table.rows[0].cells
        for cell, text in zip(hdr_cells, ["Category", "Date(s)"]):
            _set_cell_bg(cell, COLOUR_HEADER_BG)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.bold           = True
            run.font.size      = Pt(11)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        # Set column widths (left narrower, right wider)
        self._set_col_widths(table, [Inches(2.4), Inches(4.6)])

        # --- Data rows ---
        for row_idx, (category, date_entries) in enumerate(sorted(flat.items())):
            row   = table.add_row()
            cells = row.cells
            bg    = COLOUR_ROW_ALT if row_idx % 2 == 0 else COLOUR_ROW_MAIN

            # Left cell — category
            _set_cell_bg(cells[0], bg)
            _set_cell_borders(cells[0])
            cells[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p = cells[0].paragraphs[0]
            run = p.add_run(category)
            run.bold      = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

            # Right cell — dates
            _set_cell_bg(cells[1], bg)
            _set_cell_borders(cells[1])
            cells[1].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            right_para = cells[1].paragraphs[0]

            for i, (d, label, resolved) in enumerate(date_entries):
                if i > 0:
                    right_para = cells[1].add_paragraph()

                if resolved:
                    date_str = d.strftime("%d %B %Y")
                    run      = right_para.add_run(f"• {date_str}")
                    run.font.size = Pt(10)
                    # Add raw/label as italic note if it differs meaningfully
                    if label and label.lower() not in date_str.lower():
                        note = right_para.add_run(f"  ({label})")
                        note.italic    = True
                        note.font.size = Pt(9)
                        note.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
                else:
                    run = right_para.add_run(f"⚠ {label}")
                    run.font.size      = Pt(10)
                    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
                    run.italic         = True

        doc.add_paragraph()  # spacer after table

    # ------------------------------------------------------------------
    # Branch-tree appendix
    # ------------------------------------------------------------------

    def _write_branch_appendix(self, doc: Document, tree: BranchTree) -> None:
        doc.add_page_break()
        hdr = doc.add_paragraph()
        run = hdr.add_run("Appendix — Branched Timeline")
        run.bold           = True
        run.font.size      = Pt(14)
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

        note = doc.add_paragraph()
        note.add_run(
            "Relative dates are shown indented beneath their anchor. "
            "Branches arise when an anchor date has multiple resolutions."
        ).font.size = Pt(9)

        tree_text = tree.render_text()
        para = doc.add_paragraph()
        run  = para.add_run(tree_text)
        run.font.name = "Courier New"
        run.font.size = Pt(8)

    # ------------------------------------------------------------------
    # Utility
    # ------------------------------------------------------------------

    @staticmethod
    def _set_col_widths(table, widths) -> None:
        for row in table.rows:
            for cell, width in zip(row.cells, widths):
                cell.width = width
