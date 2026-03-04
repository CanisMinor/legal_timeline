"""
analyser.py — high-level API: feed in a .docx path, get back an AnalysisResult.

Usage::

    from legal_timeline import DocumentAnalyser

    analyser = DocumentAnalyser()
    result   = analyser.analyse("shareholders_agreement.docx")

    print(result.branch_tree())          # ASCII tree
    result.export_docx("timeline.docx")  # Word report
"""

from __future__ import annotations

import logging
from dataclasses import dataclass, field
from pathlib     import Path
from typing      import List, Optional

from docx import Document as DocxDocument

from .models      import CategorisedDate, DateEntry
from .extractor   import DateExtractor
from .categoriser import CategoryRule, DateCategoriser, DEFAULT_RULES
from .timeline    import BranchTree
from .reporter    import TimelineReporter

log = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Result container
# ---------------------------------------------------------------------------

@dataclass
class AnalysisResult:
    """
    Everything produced by ``DocumentAnalyser.analyse()``.

    Attributes
    ----------
    source_path:    Path to the input document.
    raw_entries:    Raw ``DateEntry`` objects before categorisation.
    categorised:    ``CategorisedDate`` objects after classification.
    tree:           The ``BranchTree`` built from the categorised entries.
    """
    source_path:  Path
    raw_entries:  List[DateEntry]
    categorised:  List[CategorisedDate]
    tree:         BranchTree

    def branch_tree(self) -> str:
        """Return the ASCII branch-tree as a string."""
        return self.tree.render_text()

    def flat_summary(self) -> str:
        """
        Return a compact, human-readable summary grouped by category.
        """
        lines = ["KEY DATES SUMMARY", "=" * 60]
        for category, entries in sorted(self.tree.flat_view().items()):
            lines.append(f"\n{category}")
            lines.append("-" * len(category))
            for d, label, resolved in entries:
                if resolved:
                    lines.append(f"  {d.strftime('%d %b %Y')}  ({label})")
                else:
                    lines.append(f"  ⚠  {label}")
        return "\n".join(lines)

    def export_docx(
        self,
        output_path: str | Path,
        title: str = "Key Dates — Legal Timeline",
        show_branch_tree: bool = True,
    ) -> Path:
        """
        Write a formatted Word document to *output_path*.

        Returns the resolved output ``Path``.
        """
        reporter = TimelineReporter(title=title, show_branch_tree=show_branch_tree)
        return reporter.generate(
            tree=self.tree,
            output_path=output_path,
            source_name=self.source_path.name,
        )


# ---------------------------------------------------------------------------
# Analyser
# ---------------------------------------------------------------------------

class DocumentAnalyser:
    """
    Extract and categorise key dates from a Word document.

    Parameters
    ----------
    dayfirst:
        How to interpret ambiguous numeric dates (e.g. ``01/02/2025``).
        ``True`` → dd/mm (European, default).
        ``False`` → mm/dd (American).
    extra_rules:
        Additional ``CategoryRule`` objects prepended to the default ruleset.
        Use this to add domain-specific categories without replacing the
        built-in ones.
    custom_rules:
        If provided, *replaces* the default ruleset entirely.
    """

    def __init__(
        self,
        dayfirst:     bool                     = True,
        extra_rules:  Optional[List[CategoryRule]] = None,
        custom_rules: Optional[List[CategoryRule]] = None,
    ) -> None:
        self.dayfirst = dayfirst

        rules = custom_rules if custom_rules is not None else list(DEFAULT_RULES)
        if extra_rules:
            rules = list(extra_rules) + rules

        self._extractor   = DateExtractor(dayfirst=dayfirst)
        self._categoriser = DateCategoriser(rules=rules)

    # ------------------------------------------------------------------
    # Main entry point
    # ------------------------------------------------------------------

    def analyse(self, docx_path: str | Path) -> AnalysisResult:
        """
        Read a ``.docx`` file and return an :class:`AnalysisResult`.

        Parameters
        ----------
        docx_path:
            Path to the Word document to analyse.
        """
        docx_path = Path(docx_path)
        log.info("Analysing %s", docx_path)

        text        = self._extract_text(docx_path)
        raw_entries = self._extractor.extract(text)

        log.info("Extracted %d date references", len(raw_entries))

        categorised = self._categoriser.categorise(raw_entries)
        tree        = BranchTree.build(categorised)

        return AnalysisResult(
            source_path=docx_path,
            raw_entries=raw_entries,
            categorised=categorised,
            tree=tree,
        )

    # ------------------------------------------------------------------
    # Text extraction
    # ------------------------------------------------------------------

    @staticmethod
    def _extract_text(docx_path: Path) -> str:
        """
        Extract plain text from a ``.docx`` file.

        Reads body paragraphs and table cells (contracts often embed dates
        in tables).
        """
        doc    = DocxDocument(str(docx_path))
        parts: List[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        parts.append(text)

        return "\n".join(parts)
